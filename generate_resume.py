#!/usr/bin/env python3
"""Generate an ATS-friendly one-page DOCX resume for Farhan Tahmid."""

from __future__ import annotations

import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUTPUT = "Farhan_Tahmid_Resume.docx"
FONT = "Calibri"


def set_run_font(run, size_pt: float, bold: bool = False, italic: bool = False) -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic


def configure_paragraph(
    paragraph,
    *,
    space_before: float = 0,
    space_after: float = 0,
    line_spacing: float = 1.0,
    left_indent: float | None = None,
) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if left_indent is not None:
        pf.left_indent = Inches(left_indent)
        pf.first_line_indent = Inches(-0.12)


def add_text_paragraph(
    doc: Document,
    text: str,
    *,
    size: float,
    bold: bool = False,
    italic: bool = False,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    space_before: float = 0,
    space_after: float = 0,
    line_spacing: float = 1.0,
):
    p = doc.add_paragraph()
    p.alignment = align
    configure_paragraph(
        p,
        space_before=space_before,
        space_after=space_after,
        line_spacing=line_spacing,
    )
    run = p.add_run(text)
    set_run_font(run, size, bold=bold, italic=italic)
    return p


def add_bullet(
    doc: Document,
    parts: list[tuple[str, bool, bool]] | str,
    *,
    size: float,
    space_after: float,
    line_spacing: float,
):
    """Standard Word-visible bullet via • character (ATS-safe plain text)."""
    p = doc.add_paragraph()
    configure_paragraph(
        p,
        space_before=0,
        space_after=space_after,
        line_spacing=line_spacing,
        left_indent=0.15,
    )
    bullet_run = p.add_run("• ")
    set_run_font(bullet_run, size)
    if isinstance(parts, str):
        parts = [(parts, False, False)]
    for text, bold, italic in parts:
        run = p.add_run(text)
        set_run_font(run, size, bold=bold, italic=italic)
    return p


def add_section_header(doc: Document, text: str, size: float, line_spacing: float, space_before: float) -> None:
    p = add_text_paragraph(
        doc,
        text,
        size=size,
        bold=True,
        space_before=space_before,
        space_after=6,
        line_spacing=line_spacing,
    )
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def build_resume(
    *,
    body_size: float,
    line_spacing: float,
    section_before: float,
    bullet_after: float,
) -> Document:
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    add_text_paragraph(
        doc,
        "Farhan Tahmid",
        size=14,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
        line_spacing=1.0,
    )
    add_text_paragraph(
        doc,
        "spacecowboy98@gmail.com | (469) 912-5790 | github.com/tahmidft | Irving, TX | U.S. Citizen",
        size=body_size,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=0,
        line_spacing=1.0,
    )
    add_text_paragraph(
        doc,
        "Portfolio: farhantahmid.vercel.app",
        size=body_size,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
        line_spacing=1.0,
    )

    add_section_header(doc, "PROFESSIONAL EXPERIENCE", 12, line_spacing, section_before)

    add_text_paragraph(
        doc,
        "Amazon Web Services | Cloud Support Engineer I | Feb 2023 – Feb 2025",
        size=body_size,
        bold=True,
        space_before=1,
        space_after=4,
        line_spacing=line_spacing,
    )

    add_text_paragraph(
        doc,
        "Internal SDE Internship",
        size=body_size,
        italic=True,
        space_after=2,
        line_spacing=line_spacing,
    )
    for bullet in [
        "Built a RAG+LLM knowledge extraction pipeline using AWS Bedrock, LangChain, and Amazon Neptune; ingested 1,000+ documents, reducing manual review time by 60% and improving retrieval accuracy by 70%.",
        "Designed a Lambda-based API service with Python, Smithy IDL, and S3, implementing full error handling, throttling, and request validation for concurrent production workloads.",
        "Implemented entity disambiguation logic using Bedrock reasoning to deduplicate LLM-extracted entities against Neptune, enabling human-in-the-loop review for novel entries.",
        "Developed infrastructure as code with CDK (TypeScript), cutting deployment times from hours to minutes and resolving a critical defect in Smithy IDL tooling.",
    ]:
        add_bullet(doc, bullet, size=body_size, space_after=bullet_after, line_spacing=line_spacing)

    add_text_paragraph(
        doc,
        "Cloud Support Engineering",
        size=body_size,
        italic=True,
        space_before=2,
        space_after=2,
        line_spacing=line_spacing,
    )
    for bullet in [
        "Resolved 500+ technical support cases for enterprise customers across Lambda, API Gateway, Cognito, and AppSync, reducing incident response time by 80%.",
        "Diagnosed complex serverless failures (Lambda cold starts, VPC timeouts, IAM permission boundaries, Step Functions) under SLA-bound response windows.",
        "Managed Sev1–Sev5 production incidents, authored root cause analyses, and delivered post-incident reports to internal engineering teams to drive service improvements.",
    ]:
        add_bullet(doc, bullet, size=body_size, space_after=bullet_after, line_spacing=line_spacing)

    add_text_paragraph(
        doc,
        "Freelance Full-Stack Engineer | Mar 2025 – Present",
        size=body_size,
        bold=True,
        space_before=3,
        space_after=4,
        line_spacing=line_spacing,
    )
    for bullet in [
        "Developed backend-driven web applications with modular architecture and automated testing, reducing bug recurrence by 40% and accelerating feature delivery by 25%.",
        "Architected a full-featured board game companion app from database schema through frontend delivery, with social sharing and game history tracking – boosting engagement by 60%, cutting manual entry errors by 80%, and reducing lookup time from minutes to seconds.",
        "Created a real-time backend with custom APIs and external data processing, improving data sync efficiency by 70% and powering a daily health-check dashboard.",
    ]:
        add_bullet(doc, bullet, size=body_size, space_after=bullet_after, line_spacing=line_spacing)

    add_section_header(doc, "PERSONAL PROJECTS", 12, line_spacing, section_before)

    projects = [
        (
            "ClearClause",
            " – AI contract analyzer (Gemini + FastAPI, Supabase) summarizing clauses, providing preference-weighted guidance, and detecting scam risks. [React, TypeScript, FastAPI, Gemini, Supabase, Render]",
        ),
        (
            "Nexus PM",
            " – Multi-tenant project workspace with RBAC, Kanban, WBS, critical-path scheduling, earned-value metrics, and full audit trail. [NestJS, Angular, TypeORM, Neon, Docker, JWT]",
        ),
        (
            "COUR",
            " – Anime tracker with custom search ranking, personalized recommendations, and cron-driven email alerts on top of AniList. [React, Vite, Vercel, Supabase, AniList GraphQL, Resend]",
        ),
        (
            "Flame Sentinel",
            " – Real-time fire alarm detector on Raspberry Pi Zero W using C++ FFT pipeline (ALSA, FFTW3), achieving 98% accuracy and <500ms latency; Flask API + React dashboard. [C++17, FFTW3, ALSA, Flask, React, Raspberry Pi]",
        ),
        (
            "Enterprise RBAC Task Manager",
            " – Multi-tenant NX monorepo (NestJS API + Angular 18) with 3-tier RBAC, JWT auth, and full audit logging. [NestJS, Angular 18, TypeScript, TypeORM, SQLite, JWT, NX]",
        ),
    ]
    for title, rest in projects:
        add_bullet(
            doc,
            [(title, True, False), (rest, False, False)],
            size=body_size,
            space_after=bullet_after,
            line_spacing=line_spacing,
        )

    add_section_header(doc, "TECHNICAL SKILLS", 12, line_spacing, section_before)
    for bullet in [
        "Languages: Python, Java, TypeScript, C++, SQL",
        "Frameworks & Tools: NestJS, Angular, React, Flask, FastAPI, LangChain, Smithy IDL, Vite, Tailwind, TanStack Query",
        "Cloud (AWS): Lambda, Bedrock, Neptune, CDK, EC2, S3, RDS, Cognito, Connect, Amplify, CloudWatch",
        "Other: Node.js, Docker, Kubernetes, Git, Postman, Linux, Raspberry Pi, Gemini API, Supabase, Neon, Resend",
    ]:
        add_bullet(doc, bullet, size=body_size, space_after=bullet_after, line_spacing=line_spacing)

    add_section_header(doc, "EDUCATION", 12, line_spacing, section_before)
    for bullet in [
        "M.S. in Computer Science – Georgia Institute of Technology (OMSCS) – AI Specialization",
        "B.S. in Computer Science – University of Texas at Dallas, Richardson, TX",
    ]:
        add_bullet(doc, bullet, size=body_size, space_after=bullet_after, line_spacing=line_spacing)

    return doc


def pdf_page_count(pdf_path: Path) -> int:
    result = subprocess.run(
        ["pdfinfo", str(pdf_path)],
        check=True,
        capture_output=True,
        text=True,
    )
    for line in result.stdout.splitlines():
        if line.startswith("Pages:"):
            return int(line.split(":")[1].strip())
    raise RuntimeError("Could not read page count from pdfinfo")


def convert_to_pdf(docx_path: Path, out_dir: Path) -> Path:
    subprocess.run(
        [
            "libreoffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(out_dir),
            str(docx_path),
        ],
        check=True,
        capture_output=True,
    )
    return out_dir / f"{docx_path.stem}.pdf"


def main() -> None:
    # Prefer requested sizes; tighten only if needed to fit one page.
    candidates = [
        # body, line_spacing, section_before, bullet_after
        (11.0, 1.05, 6, 1.5),
        (11.0, 1.0, 5, 1.0),
        (10.5, 1.02, 5, 1.0),
        (10.5, 1.0, 4, 0.5),
        (10.5, 0.98, 3, 0.5),
        (10.5, 0.95, 2, 0.25),
        (10.5, 0.92, 2, 0.0),
        (10.0, 0.98, 3, 0.5),
        (10.0, 0.95, 2, 0.25),
    ]

    out = Path(OUTPUT)
    chosen = None
    pages = None

    for body_size, line_spacing, section_before, bullet_after in candidates:
        doc = build_resume(
            body_size=body_size,
            line_spacing=line_spacing,
            section_before=section_before,
            bullet_after=bullet_after,
        )
        doc.save(out)
        pdf = convert_to_pdf(out, out.parent)
        pages = pdf_page_count(pdf)
        print(
            f"try body={body_size} leading={line_spacing} "
            f"section_before={section_before} bullet_after={bullet_after} -> {pages} page(s)"
        )
        if pages == 1:
            chosen = (body_size, line_spacing, section_before, bullet_after)
            break

    if chosen is None:
        print("WARNING: could not fit onto one page with available candidates", file=sys.stderr)
        sys.exit(1)

    print(f"Saved {OUTPUT} and {out.with_suffix('.pdf').name} ({pages} page, body={chosen[0]}pt)")


if __name__ == "__main__":
    main()
